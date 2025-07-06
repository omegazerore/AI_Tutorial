import logging
from docx import Document
from docx.shared import Pt
from typing import List, Dict

from src.logic.trend_week_report.utils.reference import reference_cleansing  # assumes you extract that to a utils module

logger = logging.getLogger(__name__)


def apply_custom_styles(document: Document) -> None:
    """Applies custom font and size styles to the Word document.

    This function modifies the default 'Normal', 'Body Text', and heading
    styles (Heading 1, 2, 3) to use Calibri and Calibri Light with specific font sizes.

    Args:
        document: A python-docx Document object.
    """
    normal = document.styles['Normal'].font
    normal.name = 'Calibri'
    normal.size = Pt(11)

    for heading, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        font = document.styles[heading].font
        font.name = 'Calibri Light'
        font.size = Pt(size)

    body = document.styles['Body Text'].font
    body.name = 'Calibri'
    body.size = Pt(11)


def add_paragraph(doc: Document, text: str, style: str = 'Body Text', space_after: Pt = Pt(6)):
    """Adds a paragraph with optional style and spacing.

    Args:
        doc: The Word document object.
        text: The text content of the paragraph.
        style: The style name to be applied to the paragraph.
        space_after: The space after the paragraph (default is 6pt).

    Returns:
        The created paragraph object.
    """
    para = doc.add_paragraph(text, style=style)
    para.paragraph_format.space_after = space_after
    return para


def add_bullet_list(doc: Document, items: List[str]):
    """Adds a bullet list to the document.

    Each item is added as a bullet point with font size set to 11pt.

    Args:
        doc: The Word document object.
        items: A list of strings to be displayed as bullet items.
    """
    for item in items:
        para = doc.add_paragraph(style='List Bullet')
        run = para.add_run(item)
        run.font.size = Pt(11)


def add_recommendations_table(doc: Document, recommendations: List):
    """Adds a recommendations table with three columns to the document.

    The table includes headers: Focus Area, Portfolio Action, and Regional Opportunity.

    Args:
        doc: The Word document object.
        recommendations: A list of objects with attributes: focus_area,
            portfolio_action, and regional_opportunity.
    """
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Focus Area'
    hdr_cells[1].text = 'Portfolio Action'
    hdr_cells[2].text = 'Regional Opportunity'

    for rec in recommendations:
        row = table.add_row().cells
        row[0].text = rec.focus_area
        row[1].text = rec.portfolio_action
        row[2].text = rec.regional_opportunity


def generate_trend_week_report(
    megatrend_context: Dict[str, str],
    megatrend_summaries: Dict,
    subtrend_2_reference: Dict[str, List[str]],
    doc_index_map: Dict[str, int],
    summary
) -> Document:
    """Generates a strategy report as a Word document.

    The report includes megatrend findings, action points, summary trends,
    recommendations, and a bibliography.

    Args:
        megatrend_context: Dictionary mapping megatrend names to context descriptions.
        megatrend_summaries: Dictionary mapping megatrends to summary objects.
            Each summary should have `.insight`, `.key_findings`, and `.result`.
        subtrend_2_reference: Dictionary mapping key finding names to reference document filenames.
        doc_index_map: Dictionary mapping document filenames to bibliography indices.
        summary: An object with `.trends`, each having `.name`, `.explanation`,
            and `.recommendations`.

    Returns:
        A Word Document object containing the formatted report.
    """

    doc = Document()
    apply_custom_styles(doc)

    doc.add_heading('Strategy Input Report', level=0)

    for megatrend in megatrend_context:
        doc.add_heading(megatrend, level=1)

        doc.add_heading("Key Findings", level=2)
        add_paragraph(doc, megatrend_summaries[megatrend].insight)

        findings = []
        for kf in megatrend_summaries[megatrend].key_findings:
            refs = [doc_index_map[f] for f in subtrend_2_reference.get(kf.name, [])]
            cleaned = reference_cleansing(sorted(set(refs)))
            findings.append(f"{kf.name} {cleaned}")
        add_bullet_list(doc, findings)

        doc.add_heading("Action Points", level=2)
        add_bullet_list(doc, [a.name for a in megatrend_summaries[megatrend].result])

    doc.add_heading("Summary", level=1)
    for idx, trend in enumerate(summary.trends):
        doc.add_heading(f"{idx + 1}. {trend.name}", level=2)
        add_paragraph(doc, trend.explanation)
        doc.add_heading("Recommendations", level=3)
        add_recommendations_table(doc, trend.recommendations)

    doc.add_heading("Bibliography", level=1)
    for filename, index in sorted(doc_index_map.items(), key=lambda x: x[1]):
        add_paragraph(doc, f"[{index}]: {filename}")

    return doc