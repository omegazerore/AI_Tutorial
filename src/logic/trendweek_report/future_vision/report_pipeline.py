"""
Hi Meng-Chieh,

Since I’m here waiting anyways, I just had a look at your latest outcome 😃 Looks very promising and I would like to now do the final test of feeding all 194
reports into your solution, then compare your output with the actual Future Vision presentation I made back in November. That would really show us how the
results compare, if one is better or if they are identical.

You find all reports that I read and combed through in November here:
https://cloud.cosnova.com/index.php/s/T0mJzDZyobijA77

Regarding the web-supported research for examples, there were only few subtrends that have no concrete examples in your latest output
(Green Innovation, AI Integration), which is already a great outcome. It’s much less work to manually research examples for 2 topics,
versus having to do so for ALL topics. Since you’re saying web search is not implemented YET, I suppose it’s still possible though?

Regarding the level of details, can you expand the description to 2-3 sentences so we get more context about each subtrend.
And for the examples, in most cases only the brand is named but not the specific service/campaign/product.
Is there an option to name both the brand and the matching product? Otherwise I would have to do another step, asking Google/GPT to find me the matching product.
"""
"""
Trend Week Report Generator

This module processes raw trend notes to generate a structured future vision report,
including megatrend aggregation, reference matching, and enhanced subtrend examples
via web search. The output is a formatted Word (.docx) document.

Author: Meng-Chieh Ling
Date: [Auto-updated]
"""
import json
import logging
from collections import defaultdict
from typing import List

from docx import Document
from tqdm import tqdm

from src.initialization import model_activation
from src.logic.trendweek_report.future_vision import megatrend_aggregation_pipeline, reference_matching_pipeline
from src.logic.trendweek_report.future_vision.utils_websearch import examples_with_deep_websearch
from src.logic.trendweek_report.word_tool import add_indented_paragraph

logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)


def reference_cleansing(references: List[int]):
    """Converts a list of integers into a condensed list of reference ranges.

    For example: [1, 2, 3, 5] -> ['1-3', '5']

    Args:
        references: A list of reference integers.

    Returns:
        A list of strings representing condensed references.
    """
    if not references:
        return []

    references.sort()
    previous_idx = references[0]
    cleaned_references = [previous_idx]

    for idx in references[1:]:

        if idx - previous_idx == 1:
            previous_idx = idx
        else:
            if previous_idx - cleaned_references[-1] > 1:
                cleaned_references[-1] = f"{cleaned_references[-1]}-{previous_idx}"
            else:
                cleaned_references[-1] = str(previous_idx)

            cleaned_references.append(idx)
            previous_idx = idx

    if previous_idx - cleaned_references[-1] > 1:
        cleaned_references[-1] = f"{cleaned_references[-1]}-{previous_idx}"
    else:
        cleaned_references[-1] = str(previous_idx)

    return cleaned_references


class TrendWeekReport:
    """Generates a Future Vision report from structured trend notes.

    Attributes:
        report_json_file: Path to the raw JSON file of trend notes.
        model: Activated LLM model used for summarization and matching.
    """
    def __init__(self, report_json_file: str, model: str):
        """Initializes the TrendWeekReport instance.

        Loads the trend notes and initializes pipelines for aggregation and matching.

        Args:
            report_json_file: Path to the JSON file containing trend notes.
            model: The name of the LLM model to use.
        """
        # load the json files directly
        self.report_json_file = report_json_file

        self.model = model_activation(model_name=model)

        self._load_notes()
        self._megatrend_aggregation_pipeline = self._build_megatrend_aggregation_pipeline()
        self._reference_matching_pipeline = self._build_reference_matching_pipeline()

    def _load_notes(self):
        """Loads trend notes from the specified JSON file.

        Raises:
            FileNotFoundError: If the file path is invalid.
            JSONDecodeError: If the JSON is malformed.
        """
        logger.info(f"Start loading the {self.report_json_file}")
        try:
            with open(self.report_json_file, 'r') as f:
                self._notes = json.load(f)
        except FileNotFoundError:
            logger.error(f"File not found: {self.report_json_file}")
            raise
        except json.JSONDecodeError:
            logger.error("Invalid JSON format")
            raise

    def _build_megatrend_aggregation_pipeline(self):
        """Builds the megatrend aggregation pipeline.

       Returns:
           A composed pipeline for megatrend summarization.
       """
        prompt_template = megatrend_aggregation_pipeline.build_prompt_template()

        return  prompt_template | self.model | megatrend_aggregation_pipeline.output_parser

    def _build_reference_matching_pipeline(self):
        """Builds the reference matching pipeline.

        Returns:
            A composed pipeline for matching subtrends to references.
        """
        prompt_template = reference_matching_pipeline.build_prompt_template()

        return prompt_template | self.model | reference_matching_pipeline.output_parser

    def _megatrend_aggregation(self):
        """Aggregates megatrend content from notes into a structured dictionary.

        Iterates over all trend notes and organizes them into self._megatrend_dict.
        """
        logger.info("\n\nStart megatrend aggregation\n\n")
        self._megatrend_dict = {}

        for filename, note in tqdm(self._notes.items()):
            logger.info(f"Processing file: {filename}")
            trends = note.get('trends', [])

            for trend in trends:
                megatrend = trend.get('name')
                subtrend_dict = self._extract_subtrend_dict(trend.get('content', []))

                self._merge_into_megatrend_dict(megatrend, subtrend_dict)

        logger.info(f"All the megatrends: {self._megatrend_dict.keys()}")

    def _extract_subtrend_dict(self, subtrends: List[dict]) -> dict:
        """Extracts subtrend information from trend content into a dictionary.

        Args:
            subtrends: A list of subtrend dictionaries.

        Returns:
            A dictionary mapping subtrend names to their definition and examples.
        """
        return {
            subtrend['name']: {
                'Definition': subtrend['definition'],
                'Examples': subtrend['examples']
            }
            for subtrend in subtrends
        }

    def _merge_into_megatrend_dict(self, megatrend: str, subtrend_dict: dict):
        """Merges new subtrend content into the existing megatrend dictionary.

        Args:
            megatrend: The name of the megatrend.
            subtrend_dict: A dictionary of subtrend data to merge.
        """
        if megatrend not in self._megatrend_dict:
            self._megatrend_dict[megatrend] = subtrend_dict
            return

        for subtrend, values in subtrend_dict.items():
            if subtrend in self._megatrend_dict[megatrend]:
                self._megatrend_dict[megatrend][subtrend]['Definition'] += f"\n\n{values['Definition']}"
                self._megatrend_dict[megatrend][subtrend]['Examples'] += f"\n\n{values['Examples']}"
            else:
                self._megatrend_dict[megatrend][subtrend] = values

    def run(self, search_context_size: str, search_model: str):
        """Executes the full report generation pipeline.

        Args:
            search_context_size: Size of the context for the web search model.
            search_model: The model to use for web search enhancement.

        Returns:
            A finalized Word document containing the Future Vision report.
        """
        self._megatrend_aggregation() # Aggregate the megatrend content from various sourc

        megatrend_context = self._prepare_megatrend_context()
        megatrend_summaries = self._summarize_megatrends(megatrend_context)
        subtrend_2_reference, doc_index_map, summary_batches = self._match_references(megatrend_summaries)
        examples_hash_map = self._enhance_examples(megatrend_summaries, search_context_size, search_model)

        final_doc = self._build_final_document(
            megatrend_summaries, subtrend_2_reference, doc_index_map, examples_hash_map
        )
        return final_doc

    def _prepare_megatrend_context(self) -> dict:
        """Formats megatrend and subtrend content into context strings.

        Returns:
            A dictionary mapping each megatrend to a formatted context string.
        """
        megatrend_context = {}
        for megatrend, subtrend_dict in self._megatrend_dict.items():
            context = ""
            for subtrend, content in subtrend_dict.items():
                context += f"**{subtrend}**\n - Definition: {content['Definition']}\n - Examples: {content['Examples']}\n\n"
            megatrend_context[megatrend] = context
        return megatrend_context

    def _summarize_megatrends(self, megatrend_context: dict) -> dict:
        """Summarizes each megatrend using the aggregation pipeline.

        Args:
            megatrend_context: Dictionary of formatted megatrend content.

        Returns:
            A dictionary mapping megatrends to summarized output.
        """
        summaries = {}
        for megatrend, text in tqdm(megatrend_context.items()):
            logger.info(f"Megatrend: {megatrend} fusion")
            result = self._megatrend_aggregation_pipeline.invoke({"text": text})
            summaries[megatrend] = result
        return summaries

    def _match_references(self, summaries: dict):
        """Matches subtrends to source documents using reference matching.

        Args:
            summaries: Dictionary of summarized megatrend outputs.

        Returns:
            subtrend_2_reference: Mapping of subtrend names to reference filenames.
            doc_index_map: Mapping of filenames to bibliography indices.
            summary_batches: Data used for reference matching batch processing.
        """
        subtrend_2_reference = defaultdict(list)
        doc_index_map = {}
        reference_index = 0
        summary_batches = []

        for megatrend, summary in summaries.items():
            for note_idx, (filename, note) in enumerate(self._notes.items()):
                if megatrend not in [t['name'] for t in note['trends']]:
                    continue

                subtrend_text = ""
                for trend in note['trends']:
                    if trend['name'] == megatrend:
                        for subtrend in trend['content']:
                            subtrend_text += f"**{subtrend['name']}**: \n - Definition: {subtrend['definition']}\n - Examples: {subtrend['examples']}\n\n"

                for subtrend in summary.result:
                    row = {
                        "note": subtrend_text,
                        "subtrend": f"**{subtrend.name}**:\n - Definition: {subtrend.definition}\n - Examples: {subtrend.examples}",
                        "note_idx": note_idx,
                        "subtrend_name": subtrend.name,
                        "filename": filename
                    }
                    summary_batches.append(row)

        outputs = self._reference_matching_pipeline.batch(summary_batches)

        for output, row in zip(outputs, summary_batches):
            if output.result == 'YES':
                name = row['filename']
                subtrend_2_reference[row['subtrend_name']].append(name)
                if name not in doc_index_map:
                    doc_index_map[name] = reference_index
                    reference_index += 1

        return subtrend_2_reference, doc_index_map, summary_batches

    def _enhance_examples(self, summaries: dict, search_context_size: str, search_model: str) -> dict:
        """Enhances subtrend examples using a deep web search service.

        Args:
            summaries: Dictionary of summarized megatrends.
            search_context_size: Context window size for search.
            search_model: Model name used for search-based enhancement.

        Returns:
            A dictionary mapping (megatrend, subtrend) pairs to enhanced examples.
        """
        batches = []
        for megatrend, summary in summaries.items():
            for subtrend in summary.result:
                batches.append({
                    "description": subtrend.definition,
                    "example": subtrend.examples,
                    "megatrend": megatrend,
                    "subtrend": subtrend.name
                })

        outputs = examples_with_deep_websearch.service(batches, search_context_size, search_model)
        examples_map = {}

        for batch, result in zip(batches, outputs):
            enhanced = ''
            for example in result['final_output'].result:
                urls = "".join(f"        ** {url.name}\n" for url in example.urls)
                enhanced += f"  - {example.brand}\n    * Action: {example.content}\n    * Source:\n{urls}"
            key = (batch['megatrend'], batch['subtrend'])
            examples_map[key] = (enhanced, batch['example'])

        return examples_map

    def _build_final_document(self, summaries, references, doc_index_map, examples_map) -> Document:
        """Constructs a well-formatted Word document for the Future Vision Report.

        Args:
            summaries: Megatrend summaries from the model.
            references: Mapping of subtrends to their references.
            doc_index_map: Mapping of filenames to bibliography index.
            examples_map: Enhanced examples for each subtrend.

        Returns:
            A formatted Word document.
        """
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        reverse_map = {v: k for k, v in doc_index_map.items()}
        doc = Document()

        # Title: Centered, large font
        title = doc.add_heading('Future Vision Report', level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for megatrend, summary in summaries.items():
            doc.add_page_break()
            doc.add_heading(megatrend, level=2)

            for idx, subtrend in enumerate(summary.result, start=1):
                ref_indices = [doc_index_map[name] for name in references[subtrend.name]]
                ref_cleaned = reference_cleansing(ref_indices)

                # Subtrend name and references
                para = doc.add_paragraph()
                run = para.add_run(f"{idx}. {subtrend.name} [{', '.join(ref_cleaned)}]")
                run.bold = True

                # Definition
                para_def = doc.add_paragraph()
                run_def = para_def.add_run("Definition: ")
                run_def.bold = True
                para_def.add_run(subtrend.definition)

                # Examples
                examples, original = examples_map[(megatrend, subtrend.name)]
                para_ex = doc.add_paragraph()
                run_ex = para_ex.add_run("Examples: ")
                run_ex.bold = True
                para_ex.add_run(original)

                # Enhanced examples as sub-bullets
                if examples.strip():
                    doc.add_paragraph("Enhanced Examples:")

                    # Split enhanced examples into distinct entries
                    for idx, example_block in enumerate(examples.strip().split("  - ")[1:], start=1):
                        lines = example_block.strip().split("\n")
                        brand_line = lines[0].strip()
                        action_line = ""
                        sources = []

                        # Parse the action and sources
                        for line in lines[1:]:
                            line = line.strip()
                            if line.startswith("* Action:"):
                                action_line = line.replace("* Action:", "").strip()
                            elif line.startswith("* Source:"):
                                continue  # header line
                            elif line.startswith("** "):
                                sources.append(line.replace("** ", "").strip())

                        # Format the enhanced example entry
                        example_para = doc.add_paragraph(f"{idx}. {brand_line}:", style='BodyText')
                        doc.add_paragraph(f"- Action: {action_line}", style='List Continue')
                        if sources:
                            if len(sources) == 1:
                                doc.add_paragraph(f"- Source: {sources[0]}", style='List Continue')
                            else:
                                doc.add_paragraph("- Sources:", style='List Continue')
                                for src in sources:
                                    doc.add_paragraph(f"{src}", style='List Bullet 2')

            doc.add_paragraph("")  # spacing

        # Bibliography section
        doc.add_page_break()
        doc.add_heading('Bibliography', level=1)
        for name, index in doc_index_map.items():
            doc.add_paragraph(f"[{index}]: {name}", style='BodyText')

        return doc

if __name__ == "__main__":

    from datetime import datetime

    from src import config

    config.debug_mode = True

    time_start = datetime.now()

    report_generator = TrendWeekReport(report_json_file="trend_week_report_raw.json",
                                       model="gpt-4.1-2025-04-14")

    report_generator.run()

    time_end = datetime.now()

    print(f"Time eclapse: {time_end - time_start}")